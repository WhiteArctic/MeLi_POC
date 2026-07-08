from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "outputs" / "documents"
OUTPUT_PATH = OUTPUT_DIR / "reporte_tecnico_solucion_poc.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 30, 45)
GRAY = RGBColor(90, 90, 90)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_YELLOW = "FFF7D6"
PALE_GREEN = "EAF5EA"
PALE_RED = "FCE8E6"
WHITE = "FFFFFF"


def set_run_font(run, size=11, bold=False, italic=False, color=INK, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx >= len(row.cells):
                continue
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D0D7DE", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_paragraph(doc, text="", style=None, size=11, bold=False, color=INK, after=6, before=0):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.10
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    paragraph.paragraph_format.space_after = Pt(8 if level == 1 else 6)
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=BLUE)
    elif level == 2:
        set_run_font(run, size=13, bold=True, color=BLUE)
    else:
        set_run_font(run, size=12, bold=True, color=DARK_BLUE)
    return paragraph


def add_bullet(doc, text, level=0):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.25 + 0.2 * level)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    return paragraph


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    return paragraph


def add_callout(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.35])
    set_table_borders(table, color="B8C7D9")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    set_run_font(run, size=11, bold=True, color=DARK_BLUE)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(body)
    set_run_font(run, size=10.5, color=INK)
    add_paragraph(doc, "", after=4)


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.35])
    set_table_borders(table, color="D0D7DE")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F8FA")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(text.strip("\n").splitlines()):
        if idx:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        set_run_font(run, size=8.5, color=RGBColor(35, 35, 35), font="Consolas")
    add_paragraph(doc, "", after=4)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, font_size=9.3):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(table)
    set_table_width(table, widths)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(header_cells[idx], header_fill)
        paragraph = header_cells[idx].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(header)
        set_run_font(run, size=font_size, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            paragraph = cells[idx].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            set_run_font(run, size=font_size, color=INK)
    set_table_width(table, widths)
    add_paragraph(doc, "", after=4)
    return table


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(1.0))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    header = section.header.paragraphs[0]
    header.text = "Reto tecnico ML - Moderacion de imagenes"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.runs[0], size=9, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.text = "Reporte tecnico de solucion POC"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(footer.runs[0], size=9, color=GRAY)
    return doc


def add_cover(doc: Document):
    add_paragraph(doc, "REPORTE TECNICO", size=10, bold=True, color=GRAY, after=4)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Documentacion tecnica de la solucion POC")
    set_run_font(run, size=24, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run(
        "Metodologia, resultados, brechas frente al reto tecnico y plan evolutivo hacia end game"
    )
    set_run_font(run, size=13, color=GRAY)
    rows = [
        ("Proyecto evaluado", "Prueba tecnica (5).zip"),
        ("Fecha del reporte", "2026-07-08"),
        ("Alcance", "POC end-to-end, datasets, golden set, evaluacion, arquitectura y roadmap"),
        ("Estado", "POC medible y extensible; requiere evolutivo para cumplir 95/95 en produccion"),
    ]
    add_table(doc, ["Campo", "Detalle"], rows, [1.65, 4.7], header_fill=LIGHT_BLUE, font_size=10)
    add_callout(
        doc,
        "Resumen de alcance",
        "La solucion desarrollada corresponde a una POC tecnica, medible y extensible. El proyecto deja implementado "
        "el flujo de datasets, controles anti-leakage, funcion de inferencia, evidencias, metricas y reportes. "
        "Como es esperable en una primera iteracion, todavia existen puntos que requieren mas analisis, evaluacion "
        "y desarrollo antes de considerarse una solucion productiva de escala.",
        fill=PALE_YELLOW,
    )
    doc.add_page_break()


def build_report():
    doc = setup_document()
    add_cover(doc)

    add_heading(doc, "1. Resumen ejecutivo", 1)
    add_paragraph(
        doc,
        "El reto tecnico solicita una POC de moderacion de imagenes que detecte textos, promociones, "
        "promesas de entrega o badges no permitidos, exponga diagnose_image(picture_url), justifique "
        "cada decision y aspire a precision y recall mayores o iguales a 95%. El proyecto revisado "
        "esta alineado con el problema y deja una base tecnica razonable para evolucionar.",
    )
    add_table(
        doc,
        ["Dimension", "Evaluacion"],
        [
            ("Alineacion con el reto", "Alta: cubre OCR, reglas, evidencia, funcion publica y evaluacion."),
            ("Madurez POC", "Buena: modular, testeable, medible y honesta sobre limitaciones."),
            ("Cumplimiento metricas", "Parcial en POC: precision candidata 0.9545 y recall 0.84 en muestra balanceada de 100 grupos."),
            ("Oportunidad principal", "Elevar recall y ampliar la validacion antes de declarar cumplimiento 95/95."),
            ("Camino recomendado", "Modelo multimodal mas fuerte, auditoria humana del golden y arquitectura productiva con colas/workers."),
        ],
        [1.9, 4.45],
        header_fill=LIGHT_BLUE,
    )

    add_heading(doc, "2. Requerimientos del PDF vs estado del proyecto", 1)
    add_table(
        doc,
        ["Requerimiento", "Estado observado", "Implicacion"],
        [
            (
                "Precision >= 95% y recall >= 95%",
                "Precision 0.9545 y recall 0.84 en candidato medido.",
                "Cumplimiento parcial; requiere evolutivo de recall antes de aceptacion final.",
            ),
            (
                "diagnose_image(picture_url) con has_infraction y evidence",
                "Implementado en detector.py con preset candidato y degradacion de dependencias.",
                "Cumple contrato funcional basico.",
            ),
            (
                "Interpretabilidad",
                "Evidencia textual, familia, scores OCR/texto/visual y estrategia.",
                "Cumple para positivos detectados; puede reforzarse en casos con OCR debil o ausente.",
            ),
            (
                "Latencia p99 <= 5000 ms",
                "p99 reportado ~2622 ms en 100 grupos con candidato; ensemble OCR supera 5000 ms.",
                "Resultado prometedor de POC; pendiente medir a escala y con concurrencia real.",
            ),
            (
                "Minimizar APIs externas de OCR",
                "Usa OCR local opcional: Tesseract/EasyOCR, sin API externa.",
                "Bien alineado con costo y privacidad.",
            ),
            (
                "Monitoreo e iteracion",
                "Documentado a nivel de bullets.",
                "Debe ampliarse a estrategia operativa con alertas, runbooks y responsables.",
            ),
            (
                "Extensibilidad",
                "Modulos policy, OCRBackend, visual classifier, datasets y evaluation desacoplados.",
                "Buen diseno para agregar nuevas politicas y modelos.",
            ),
            (
                "Infraestructura 60k RPM promedio / 100k RPM pico",
                "No hay servicio productivo con colas, workers, autoscaling o backpressure.",
                "Pendiente natural para la etapa productiva end game.",
            ),
        ],
        [1.65, 2.35, 2.35],
        header_fill=LIGHT_BLUE,
        font_size=8.6,
    )

    add_heading(doc, "3. Metodologia aplicada", 1)
    add_heading(doc, "3.1 Construccion de datasets y control anti-leakage", 2)
    add_paragraph(
        doc,
        "El proceso parte del CSV legacy con picture_url, infraction_detected, labels_detected y ocr_text. "
        "Como el sistema anterior es OCR + Regex, sus etiquetas se tratan como pseudo-etiquetas de alta "
        "confianza, no como verdad humana absoluta. Para evitar leakage, la particion usa leakage_group_id "
        "derivado del item_id parseado desde la URL.",
    )
    add_table(
        doc,
        ["Conjunto", "Uso", "Filas", "Grupos", "Puede entrenar?"],
        [
            ("train", "Entrenar modelos baseline", "377,579", "339,189", "Si"),
            ("validation", "Seleccion de umbrales y variantes", "81,255", "72,679", "No para fit final"),
            ("test_internal", "Medicion interna pre-golden", "81,373", "72,690", "No"),
            ("golden_set_v1", "Aceptacion/regresion", "5,000", "5,000", "Nunca"),
        ],
        [1.25, 2.35, 0.9, 0.9, 0.95],
        header_fill=LIGHT_BLUE,
        font_size=8.8,
    )
    add_bullet(doc, "El golden queda completamente excluido de train, validation y test_internal.")
    add_bullet(doc, "El split se hace por grupo, no por fila, para evitar imagenes hermanas en particiones distintas.")
    add_bullet(doc, "Columnas como ocr_text, labels_detected e infraction_detected no son features productivas; son metadata de entrenamiento/evaluacion.")

    add_heading(doc, "3.2 Segmentacion del golden set", 2)
    add_paragraph(
        doc,
        "El golden set se diseno balanceado para medir precision y recall sin quedar dominado por negativos. "
        "Los positivos se segmentan por familia de infraccion y los negativos por complejidad de texto OCR.",
    )
    add_table(
        doc,
        ["Tipo", "Segmentos principales", "Motivo"],
        [
            (
                "Positivos",
                "campaign_event, price_promotion, shipping_promise, marketplace_badge_social_proof, trust_payment_platform_claim, quality_originality_claim",
                "Medir recall por familia y detectar regresiones escondidas por el promedio.",
            ),
            (
                "Negativos",
                "negative_no_ocr_text, negative_short_nonpolicy_text, negative_medium_nonpolicy_text, negative_dense_nonpolicy_text",
                "Medir falsos positivos en imagenes sin texto, texto corto, texto medio y texto denso.",
            ),
            (
                "Excluidos",
                "conflictos, positivos sin etiqueta interpretable, negativos con palabras ambiguas de politica",
                "Reducir ruido de pseudo-etiquetas y sesgos del sistema legacy.",
            ),
        ],
        [1.2, 3.0, 2.15],
        header_fill=LIGHT_BLUE,
        font_size=8.5,
    )

    add_heading(doc, "3.3 Pipeline de inferencia", 2)
    add_code_block(
        doc,
        """
picture_url
  -> descarga con cache local y verify_ssl=False por entorno corporativo
  -> OCR local reemplazable: Tesseract multipass o EasyOCR
  -> normalizacion, sinonimos y correcciones OCR
  -> fuzzy matching de terminos de politica
  -> modelo textual Naive Bayes sobre texto completo y chunks
  -> senales visuales heuristicas y clasificador MobileNetV3 opcional
  -> decision hibrida con soporte/abstencion
  -> {has_infraction, evidence}
        """,
    )
    add_paragraph(
        doc,
        "La arquitectura privilegia modularidad: OCRBackend permite intercambiar Tesseract, EasyOCR, "
        "PaddleOCR o un servicio interno; policy.py concentra terminos y fuzzy matching; detector.py "
        "orquesta decision y evidencia; evaluation.py separa medicion de inferencia.",
    )

    add_heading(doc, "4. Fundamentos matematicos de evaluacion", 1)
    add_paragraph(
        doc,
        "La evaluacion se basa en matriz de confusion: TP positivos correctamente detectados, FP negativos "
        "marcados como infraccion, TN negativos correctamente rechazados y FN infracciones omitidas.",
    )
    add_table(
        doc,
        ["Metrica", "Formula", "Interpretacion para el reto"],
        [
            ("Precision", "TP / (TP + FP)", "De las bajas propuestas, cuantas son correctas. Controla falsos positivos."),
            ("Recall", "TP / (TP + FN)", "De las infracciones reales, cuantas se detectan. Controla cobertura."),
            ("F1", "2 * precision * recall / (precision + recall)", "Balance armonico entre precision y recall."),
            ("p99 latency", "percentil 99 de tiempos por imagen", "Garantiza que casi todos los casos cumplan SLA de 5000 ms."),
        ],
        [1.25, 2.0, 3.1],
        header_fill=LIGHT_BLUE,
        font_size=8.8,
    )
    add_callout(
        doc,
        "Calculo candidato medido",
        "Con TP=42, FP=2, TN=48 y FN=8: precision = 42/(42+2) = 0.9545; "
        "recall = 42/(42+8) = 0.84; F1 = 0.8936. La precision cumple el umbral, "
        "pero el recall queda 11 puntos porcentuales por debajo del objetivo de 0.95.",
        fill=PALE_YELLOW,
    )

    add_heading(doc, "5. Resultados obtenidos", 1)
    add_heading(doc, "5.1 Baselines offline", 2)
    add_table(
        doc,
        ["Baseline", "Precision", "Recall", "F1", "Lectura"],
        [
            ("Regex OCR", "1.000", "0.876", "0.934", "Muy preciso, pero pierde cobertura."),
            ("Modelo textual threshold 0.58", "0.702", "0.635", "0.667", "Aprende senales, pero no es suficiente solo."),
            ("Regex OR modelo threshold 0.99", "0.823", "0.974", "0.892", "Sube recall, pero introduce falsos positivos."),
        ],
        [1.65, 0.8, 0.8, 0.8, 2.3],
        header_fill=LIGHT_BLUE,
        font_size=8.8,
    )
    add_heading(doc, "5.2 Evaluacion end-to-end", 2)
    add_table(
        doc,
        ["Variante", "Precision", "Recall", "F1", "p99"],
        [
            ("OCR multipass + politica", "0.955", "0.420", "0.583", "4193 ms"),
            ("Normalizacion/sinonimos + policy", "0.966", "0.560", "0.709", "4411 ms"),
            ("EasyOCR opcional", "1.000", "0.680", "0.810", "2548 ms"),
            ("Tesseract + EasyOCR ensemble", "1.000", "0.740", "0.851", "6570 ms"),
            ("Preset candidato EasyOCR + MobileNetV3", "0.955", "0.840", "0.894", "2622 ms"),
        ],
        [2.65, 0.75, 0.75, 0.65, 0.85],
        header_fill=LIGHT_BLUE,
        font_size=8.5,
    )
    add_paragraph(
        doc,
        "El candidato actual mejora sustancialmente frente a reglas simples y conserva precision >= 95% "
        "en la muestra balanceada. El recall sigue siendo la limitacion principal; por tanto, la POC "
        "debe presentarse como base evolutiva y no como cumplimiento final del reto.",
    )

    add_heading(doc, "5.3 Analisis de errores", 2)
    add_table(
        doc,
        ["Tipo", "Segmentos afectados", "Causa probable", "Mejora propuesta"],
        [
            (
                "Falsos negativos",
                "marketplace_badge_social_proof, price_promotion, trust_payment_platform_claim, quality_originality_claim",
                "OCR no extrae texto, texto distorsionado o senal visual insuficiente.",
                "Modelo multimodal entrenado por familia + OCR mas robusto + hard negative/positive mining.",
            ),
            (
                "Falsos positivos",
                "negative_dense_nonpolicy_text",
                "Textos densos de libros/documentos activan modelo visual/textual.",
                "Taxonomia de excepciones por categoria, calibracion visual y dataset de negativos densos.",
            ),
            (
                "Latencia alta en ensemble",
                "OCR Tesseract + EasyOCR",
                "Multiples pasadas OCR elevan p99 por encima de SLA.",
                "Batching, cascada barata->cara, early exit y workers OCR precalentados.",
            ),
        ],
        [1.15, 1.65, 1.75, 1.8],
        header_fill=LIGHT_BLUE,
        font_size=8.2,
    )

    add_heading(doc, "6. Documentacion tecnica de la solucion", 1)
    add_table(
        doc,
        ["Componente", "Responsabilidad", "Estado POC", "Evolucion end game"],
        [
            ("image_io.py", "Descarga y cache de imagenes", "Funcional con verify_ssl=False", "Certificados corporativos configurables y politicas de cache."),
            ("ocr.py", "Backends OCR reemplazables", "Tesseract/EasyOCR/ensemble", "PaddleOCR/EasyOCR productivo con batching y warm workers."),
            ("policy.py", "Normalizacion, sinonimos y fuzzy matching", "Explicable y extensible", "Versionado de politicas y tests por familia."),
            ("text_model.py", "Clasificador textual baseline", "Naive Bayes liviano", "Modelo contextual/multimodal calibrado."),
            ("visual.py / visual_model.py", "Senales visuales y clasificador MobileNetV3", "Complementario, recall aun bajo", "SigLIP/CLIP/ConvNeXt fine-tuned multi-label."),
            ("detector.py", "Decision hibrida y evidence", "Contrato productivo implementado", "Decision engine calibrado por costo/riesgo y categorias."),
            ("evaluation.py", "Metricas globales y por segmento", "Listo para POC", "Dashboards, slice metrics y alertas operativas."),
            ("datasets.py", "Split y curado anti-leakage", "Solido", "Soporte de versiones, lineage y auditoria humana."),
        ],
        [1.25, 1.5, 1.55, 2.05],
        header_fill=LIGHT_BLUE,
        font_size=8.1,
    )

    add_heading(doc, "7. Arquitectura end game propuesta", 1)
    add_code_block(
        doc,
        """
Item/image event queue
  -> ingestion workers: validacion, deduplicacion y cache
  -> cheap filter CPU: reglas basicas, metadata, imagen corrupta
  -> OCR workers: batch OCR local, timeout, early exit
  -> visual/multimodal workers GPU: embeddings + clasificador multi-label
  -> decision service: fusion calibrada, evidence, policy version
  -> moderation result topic + storage audit
  -> review queue para baja confianza, FP/FN y muestras de monitoreo
  -> dashboards, alertas, retraining pipeline y regression suite
        """,
    )
    add_table(
        doc,
        ["Capacidad requerida", "Diseno recomendado"],
        [
            ("60k RPM promedio / 100k RPM pico", "Colas particionadas, workers horizontales, autoscaling y backpressure."),
            ("p99 <= 5000 ms", "Cascada de modelos, timeouts por etapa, early exit y batch inference."),
            ("Costos", "Evitar APIs externas; usar OCR/modelos locales; enrutar a GPU solo casos necesarios."),
            ("Robustez", "DLQ, retries idempotentes, circuit breakers, cache de imagenes y metricas por etapa."),
            ("Extensibilidad", "Politicas versionadas y salida multi-label para nuevas familias de moderacion."),
        ],
        [1.85, 4.5],
        header_fill=LIGHT_BLUE,
        font_size=8.8,
    )

    add_heading(doc, "8. Puntos de mejora y plan evolutivo", 1)
    add_table(
        doc,
        ["Punto observado", "Estado actual de la POC", "Por que importa", "Evolucion propuesta"],
        [
            (
                "Recall por debajo del objetivo",
                "Actualmente el candidato alcanza recall 0.84 en la muestra medida.",
                "El reto solicita recall >= 95%; aun hay infracciones que no se detectan.",
                "Con mas tiempo: fine-tuning visual/multimodal, salida multi-label y mining de falsos negativos.",
            ),
            (
                "Tamano de evaluacion del candidato",
                "Actualmente se midio una muestra balanceada de 100 grupos unicos para iterar rapido.",
                "Es suficiente para comparar alternativas, pero no para declarar cumplimiento estadistico.",
                "Con mas tiempo: ejecutar validation/test_internal a mayor escala y reservar golden para aceptacion final.",
            ),
            (
                "Golden set pseudo-etiquetado",
                "Actualmente el golden usa pseudo-etiquetas de alta confianza provenientes del flujo legacy.",
                "Puede heredar sesgos del sistema OCR+Regex anterior.",
                "Con mas tiempo: auditoria humana doble, adjudicacion de desacuerdos y golden_v2_human_reviewed.",
            ),
            (
                "Descarga con verify_ssl=False",
                "Actualmente se desactiva la verificacion SSL para operar desde red corporativa.",
                "Es aceptable para la POC, pero no deberia quedar como default productivo sin control.",
                "Con mas tiempo: variable de entorno, CA bundle corporativo y warning auditable.",
            ),
            (
                "Rutas locales en comandos",
                "Actualmente algunos comandos documentados apuntan al Python local usado en desarrollo.",
                "Puede dificultar la replicacion en otro equipo.",
                "Con mas tiempo: default a python, entorno virtual documentado y eliminacion de rutas absolutas.",
            ),
            (
                "Dependencias opcionales",
                "Actualmente EasyOCR, Torch y Torchvision son opcionales y el sistema degrada si faltan.",
                "La degradacion evita romper la POC, pero reduce reproducibilidad del candidato completo.",
                "Con mas tiempo: empaquetar pesos, lockfile, checksums y contenedor reproducible.",
            ),
            (
                "Arquitectura de escala",
                "Actualmente la solucion es script/function-first, adecuada para POC.",
                "El escenario productivo del reto menciona 60k RPM promedio y 100k RPM pico.",
                "Con mas tiempo: servicio con colas, workers, batching, autoscaling, DLQ y observabilidad.",
            ),
            (
                "Monitoreo operativo",
                "Actualmente se listan metricas y senales, pero no hay runbook productivo completo.",
                "Para operar el modelo se necesitan alertas, responsables y acciones definidas.",
                "Con mas tiempo: alertas por precision proxy, drift, latency p99, failure rate y review queue.",
            ),
            (
                "Documentacion de uso de IA",
                "Actualmente el uso de agentes de IA se documenta como apoyo al desarrollo.",
                "El PDF solicita explicar proposito y alcance del uso de IA.",
                "Con mas tiempo: ampliar controles, prompts relevantes, limites y validaciones humanas.",
            ),
        ],
        [1.25, 1.35, 1.55, 2.2],
        header_fill=LIGHT_BLUE,
        font_size=7.8,
    )

    add_heading(doc, "9. Estrategia de monitoreo e iteracion en produccion", 1)
    add_table(
        doc,
        ["Senal", "Alerta sugerida", "Accion"],
        [
            ("Latency p99", "> 5000 ms por 10 min o crecimiento > 30%", "Escalar workers, activar cascada barata, revisar OCR."),
            ("Failure rate descarga/OCR", "> 1% por sitio o dominio", "Reintentos, DLQ, validar cambios de CDN/red."),
            ("Positive rate por sitio/categoria", "Desvio > 3 sigma vs baseline historico", "Revisar drift o cambio de politica/campana."),
            ("Appeals/falsos positivos", "Aumento semanal > 20%", "Muestreo humano y ajuste de umbrales/reglas."),
            ("Falsos negativos de auditoria", "Recall proxy < 95%", "Agregar casos al backlog y reentrenar."),
            ("Drift OCR", "Cambios en longitud, idioma o tasa de texto vacio", "Revisar OCR/preprocesamiento y categorias afectadas."),
        ],
        [1.6, 2.15, 2.6],
        header_fill=LIGHT_BLUE,
        font_size=8.3,
    )
    add_paragraph(
        doc,
        "El ciclo de iteracion recomendado es: capturar errores productivos, priorizar por impacto, "
        "etiquetar muestras, agregar a train/calibration, validar en validation/test_internal, ejecutar "
        "regresion contra golden y desplegar solo si no degrada segmentos criticos.",
    )

    add_heading(doc, "10. Uso de agentes de IA en el desarrollo", 1)
    add_table(
        doc,
        ["Uso", "Proposito", "Control aplicado"],
        [
            (
                "Exploracion del reto y CSV",
                "Identificar columnas, distribuciones, duplicados y riesgos de leakage.",
                "Resultados verificados con scripts reproducibles.",
            ),
            (
                "Diseno del golden set",
                "Proponer segmentacion, exclusiones y reglas anti-leakage.",
                "Generacion deterministica, manifiesto y validaciones de duplicados/conflictos.",
            ),
            (
                "Scaffolding de POC",
                "Acelerar estructura modular, tests y documentacion.",
                "Tests unitarios, py_compile y reportes medibles.",
            ),
            (
                "Analisis de errores",
                "Interpretar falsos positivos/negativos y proponer mejoras.",
                "Revision contra detalles CSV y metricas reportadas.",
            ),
            (
                "Redaccion tecnica",
                "Convertir hallazgos en narrativa clara de entrega.",
                "Limitaciones declaradas explicitamente; no se inventan metricas.",
            ),
        ],
        [1.35, 2.25, 2.75],
        header_fill=LIGHT_BLUE,
        font_size=8.5,
    )
    add_callout(
        doc,
        "Limite de uso de IA",
        "Los agentes de IA se usaron como apoyo de analisis, implementacion y documentacion. "
        "No reemplazan auditoria humana del golden, aprobacion de politicas ni validacion productiva. "
        "Las decisiones finales deben respaldarse con metricas reproducibles y revision tecnica.",
        fill=PALE_YELLOW,
    )

    add_heading(doc, "11. Como ejecutar el proyecto desde cero", 1)
    add_paragraph(
        doc,
        "Los siguientes pasos permiten replicar la POC desde un entorno limpio. En red corporativa, "
        "las descargas de imagen usan verify_ssl=False por diseno de la prueba; para produccion se "
        "debe configurar certificado corporativo.",
    )
    add_numbered(doc, "Descomprimir el proyecto y ubicarse en la raiz del repositorio.")
    add_code_block(doc, "unzip 'Prueba tecnica (5).zip'\ncd 'Prueba tecnica (5)'")
    add_numbered(doc, "Crear entorno Python base e instalar dependencias minimas.")
    add_code_block(doc, "python -m venv .venv\nsource .venv/bin/activate\npython -m pip install -U pip\npython -m pip install -e .")
    add_numbered(doc, "Opcional: instalar dependencias fuertes para candidato OCR/visual.")
    add_code_block(doc, "python -m pip install -e '.[candidate]'")
    add_numbered(doc, "Validar tests unitarios.")
    add_code_block(doc, "PYTHONPATH=src python -m unittest discover -s tests -v")
    add_numbered(doc, "Regenerar datasets si se parte del CSV original y golden existente.")
    add_code_block(doc, "PYTHONPATH=src python scripts/build_datasets.py")
    add_numbered(doc, "Entrenar baseline textual.")
    add_code_block(doc, "PYTHONPATH=src python scripts/train_model.py")
    add_numbered(doc, "Evaluar baselines offline.")
    add_code_block(
        doc,
        "PYTHONPATH=src python scripts/compare_baselines.py outputs/datasets/validation.csv outputs/reports/validation_baseline_comparison.json",
    )
    add_numbered(doc, "Ejecutar smoke end-to-end desde URL.")
    add_code_block(
        doc,
        "PYTHONPATH=src python scripts/evaluate_end_to_end.py --dataset-csv outputs/datasets/validation.csv --rows-per-class 2 --output-json outputs/reports/smoke_latest.json",
    )
    add_numbered(doc, "Ejecutar candidato completo si EasyOCR/Torch estan instalados.")
    add_code_block(
        doc,
        "PYTHONPATH=src python scripts/evaluate_end_to_end.py --preset candidate --dataset-csv outputs/datasets/validation.csv --rows-per-class 50 --dedupe-group-key leakage_group_id --output-json outputs/reports/candidate_validation_latest.json",
    )
    add_numbered(doc, "Diagnosticar una imagen puntual.")
    add_code_block(doc, "PYTHONPATH=src python scripts/diagnose_url.py 'https://...'")

    add_heading(doc, "12. Conclusiones", 1)
    add_paragraph(
        doc,
        "El proyecto es una POC solida: entiende el problema, separa datasets correctamente, evita leakage, "
        "expone una funcion productiva, mide resultados y documenta puntos de mejora. Su mayor valor es que "
        "convierte un sistema legacy OCR+Regex en una plataforma evaluable y extensible.",
    )
    add_paragraph(
        doc,
        "La principal oportunidad de mejora es elevar el recall hasta el objetivo de 95%. Este resultado es "
        "normal para una POC: permite identificar donde invertir el siguiente ciclo de desarrollo. Para llegar "
        "al end game se requiere auditoria humana del golden, modelo multimodal mas fuerte, evaluacion a mayor "
        "escala y arquitectura productiva con monitoreo, colas y workers.",
    )
    add_callout(
        doc,
        "Recomendacion final",
        "Entregar el proyecto como base tecnica robusta y proponer un evolutivo enfocado en golden human-reviewed, "
        "clasificador multimodal, validacion estadistica completa y diseno productivo. Los hallazgos actuales "
        "deben leerse como insumos de priorizacion, no como fallas de la POC.",
        fill=PALE_GREEN,
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_report()
